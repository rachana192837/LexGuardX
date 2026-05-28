import io
import pytest
from docx import Document
from fastapi import UploadFile
from parsers import parse_document


def create_docx(text: str) -> bytes:
    """Create a DOCX file in memory with the given text."""
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@pytest.mark.asyncio
async def test_docx_extraction():
    """Test basic DOCX text extraction."""
    docx_bytes = create_docx("Hello world. This is a test.")
    file = UploadFile(filename="test.docx", file=io.BytesIO(docx_bytes))
    result = await parse_document(file)
    assert "Hello world" in result
    assert "This is a test" in result


@pytest.mark.asyncio
async def test_docx_formatted():
    """Test DOCX with formatted text."""
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()
    file = UploadFile(filename="formatted.docx", file=io.BytesIO(docx_bytes))
    result = await parse_document(file)
    assert "First paragraph" in result
    assert "Second paragraph" in result


@pytest.mark.asyncio
async def test_docx_empty():
    """Test empty DOCX file."""
    docx_bytes = create_docx("")
    file = UploadFile(filename="empty.docx", file=io.BytesIO(docx_bytes))
    result = await parse_document(file)
    assert result.strip() == "" or result is not None

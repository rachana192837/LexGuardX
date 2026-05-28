import io
import pytest
from fastapi import UploadFile
from httpx import AsyncClient, ASGITransport
from main import app
from agents import analyze_comparison


@pytest.mark.asyncio
async def test_analyze_comparison_returns_items():
    """Test that analyze_comparison returns risk items."""
    original = ["The company shall pay $500,000.", "Termination requires 30 days notice."]
    revised = ["The company shall pay $100,000.", "Termination requires 30 days notice."]
    result = await analyze_comparison(original, revised)
    assert isinstance(result, list)
    assert len(result) > 0
    assert "clause_id" in result[0]
    assert "change_type" in result[0]
    assert "risk_level" in result[0]
    assert "summary" in result[0]


@pytest.mark.asyncio
async def test_analyze_comparison_identical():
    """Test that identical documents return empty or low-risk items."""
    original = ["Same sentence.", "Another same sentence."]
    revised = ["Same sentence.", "Another same sentence."]
    result = await analyze_comparison(original, revised)
    assert isinstance(result, list)


@pytest.mark.asyncio
async def test_compare_endpoint_pdf():
    """Test /compare with two PDF files."""
    import fitz

    # Create simple PDFs
    def make_pdf(text):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text)
        return doc.tobytes()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        files = {
            "original_file": ("orig.pdf", make_pdf("Hello world."), "application/pdf"),
            "revised_file": ("rev.pdf", make_pdf("Hello universe."), "application/pdf"),
        }
        response = await client.post("/compare", files=files)
        assert response.status_code == 200
        data = response.json()
        assert "original" in data
        assert "revised" in data
        assert len(data["original"]) > 0
        assert len(data["revised"]) > 0


@pytest.mark.asyncio
async def test_compare_endpoint_docx():
    """Test /compare with two DOCX files."""
    from docx import Document

    def make_docx(text):
        doc = Document()
        doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        files = {
            "original_file": ("orig.docx", make_docx("Original contract terms."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "revised_file": ("rev.docx", make_docx("Revised contract terms."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        }
        response = await client.post("/compare", files=files)
        assert response.status_code == 200
        data = response.json()
        assert "original" in data
        assert "revised" in data


@pytest.mark.asyncio
async def test_compare_invalid_file_type():
    """Test /compare rejects non-PDF/DOCX files."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        files = {
            "original_file": ("test.txt", b"Hello", "text/plain"),
            "revised_file": ("test2.txt", b"World", "text/plain"),
        }
        response = await client.post("/compare", files=files)
        assert response.status_code == 400


@pytest.mark.asyncio
async def test_compare_mixed_file_types():
    """Test /compare accepts PDF + DOCX combination."""
    import fitz
    from docx import Document

    def make_pdf(text):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text)
        return doc.tobytes()

    def make_docx(text):
        doc = Document()
        doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        files = {
            "original_file": ("orig.pdf", make_pdf("PDF content."), "application/pdf"),
            "revised_file": ("rev.docx", make_docx("DOCX content."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        }
        response = await client.post("/compare", files=files)
        assert response.status_code == 200


@pytest.mark.asyncio
async def test_compare_empty_documents():
    """Test /compare handles empty documents gracefully."""
    import fitz

    def make_empty_pdf():
        doc = fitz.open()
        doc.new_page()
        return doc.tobytes()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        files = {
            "original_file": ("empty.pdf", make_empty_pdf(), "application/pdf"),
            "revised_file": ("empty2.pdf", make_empty_pdf(), "application/pdf"),
        }
        response = await client.post("/compare", files=files)
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data["original"], list)
        assert isinstance(data["revised"], list)


@pytest.mark.asyncio
async def test_compare_analyze_endpoint():
    """Test /compare/analyze endpoint."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        payload = {
            "original": ["The company shall pay $500,000."],
            "revised": ["The company shall pay $100,000."],
        }
        response = await client.post("/compare/analyze", json=payload)
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)


@pytest.mark.asyncio
async def test_compare_analyze_empty_input():
    """Test /compare/analyze with empty lists."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        payload = {
            "original": [],
            "revised": [],
        }
        response = await client.post("/compare/analyze", json=payload)
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)


@pytest.mark.asyncio
async def test_compare_analyze_identical_documents():
    """Test /compare/analyze with identical documents."""
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        payload = {
            "original": ["Same sentence.", "Another sentence."],
            "revised": ["Same sentence.", "Another sentence."],
        }
        response = await client.post("/compare/analyze", json=payload)
        assert response.status_code == 200
        data = response.json()
        assert isinstance(data, list)


@pytest.mark.asyncio
async def test_rate_limiting():
    """Test that rate limiting returns 429 when limit exceeded."""
    import fitz
    from main import _rate_limit_store

    def make_pdf(text):
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), text)
        return doc.tobytes()

    # Clear rate limit store for this test
    _rate_limit_store.clear()

    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        # Make requests until rate limited
        rate_limited = False
        for i in range(15):
            files = {
                "original_file": (f"orig{i}.pdf", make_pdf(f"Text {i}."), "application/pdf"),
                "revised_file": (f"rev{i}.pdf", make_pdf(f"Revised {i}."), "application/pdf"),
            }
            response = await client.post("/compare", files=files)
            if response.status_code == 429:
                rate_limited = True
                break
        
        assert rate_limited, "Rate limiting should have triggered within 15 requests"
